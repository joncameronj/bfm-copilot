"""
DOCX Processor - Extract text from Word documents.

Handles DOCX files (NES templates, etc.) by extracting:
- Paragraph text
- Table content
- Basic formatting preservation
"""

from pathlib import Path
from typing import Optional

from docx import Document


def extract_docx_text(docx_path: Path) -> dict:
    """
    Extract text content from a DOCX file.

    Args:
        docx_path: Path to the DOCX file

    Returns:
        Dictionary with:
        - text_content: Extracted text as markdown
        - paragraph_count: Number of paragraphs
        - table_count: Number of tables
        - has_content: Whether meaningful content was found
    """
    doc = Document(str(docx_path))

    content_parts = []
    paragraph_count = 0
    table_count = 0

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Check if it's a heading
            if para.style and para.style.name.startswith("Heading"):
                level = _get_heading_level(para.style.name)
                content_parts.append(f"{'#' * level} {text}")
            else:
                content_parts.append(text)
            paragraph_count += 1

    # Extract tables
    for table in doc.tables:
        table_md = _table_to_markdown(table)
        if table_md:
            content_parts.append(table_md)
            table_count += 1

    full_text = "\n\n".join(content_parts)

    return {
        "text_content": full_text,
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "has_content": len(full_text.strip()) > 50,
    }


def _get_heading_level(style_name: str) -> int:
    """Extract heading level from style name."""
    # Style names like "Heading 1", "Heading 2", etc.
    try:
        level = int(style_name.split()[-1])
        return min(level, 6)  # Cap at h6
    except (ValueError, IndexError):
        return 2  # Default to h2


def _table_to_markdown(table) -> Optional[str]:
    """Convert a DOCX table to markdown format."""
    rows = []

    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):  # Skip completely empty rows
            rows.append(cells)

    if not rows:
        return None

    # Build markdown table
    md_lines = []

    # Header row
    header = rows[0]
    md_lines.append("| " + " | ".join(header) + " |")
    md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")

    # Data rows
    for row in rows[1:]:
        # Pad row to match header length
        while len(row) < len(header):
            row.append("")
        md_lines.append("| " + " | ".join(row) + " |")

    return "\n".join(md_lines)


async def process_nes_template(docx_path: Path) -> dict:
    """
    Process a NES (Neuro-Energetic Scan) template document.

    These are specific clinical assessment templates.

    Args:
        docx_path: Path to the NES template DOCX

    Returns:
        Dictionary with structured content and metadata
    """
    result = extract_docx_text(docx_path)

    if not result["has_content"]:
        return {
            "text_content": "",
            "success": False,
            "error": "No content found in document",
        }

    # Add metadata header
    content = result["text_content"]
    header = f"""# NES Assessment Template

Source: {docx_path.name}
Type: Clinical Assessment Template

---

"""

    # Try to extract case study info from path
    case_study_id = _extract_case_study_id(docx_path)

    return {
        "text_content": header + content,
        "paragraph_count": result["paragraph_count"],
        "table_count": result["table_count"],
        "case_study_id": case_study_id,
        "success": True,
    }


def _extract_case_study_id(filepath: Path) -> Optional[str]:
    """Extract case study ID from file path."""
    import re

    path_str = str(filepath).lower()

    # Look for patterns like "diabetes-cs1", "thyroid-cs3", etc.
    match = re.search(r"(diabetes|thyroid|hormones|neurological|neuro)-?cs(\d+)", path_str)
    if match:
        category = match.group(1)
        if category == "neuro":
            category = "neurological"
        return f"{category}-cs{match.group(2)}"

    # Look for just "cs1", "cs2", etc. and try to infer category from parent
    match = re.search(r"cs(\d+)", path_str)
    if match:
        # Try to find category in path
        for cat in ["diabetes", "thyroid", "hormones", "neurological"]:
            if cat in path_str:
                return f"{cat}-cs{match.group(1)}"

    return None
